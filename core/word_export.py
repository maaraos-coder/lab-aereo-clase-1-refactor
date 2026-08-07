"""Generación de documentos Word editables para laboratorios y evaluaciones.

El módulo no depende de Streamlit. Recibe datos ya cargados por la aplicación y
retorna bytes de archivos ``.docx`` o ``.zip`` listos para descargar.
"""
from __future__ import annotations

import ast
import io
import json
import re
import zipfile
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


COURSE_TITLE = "Curso 1 · Aislamiento acústico al ruido aéreo"
LAB_TITLES = {
    1: "Laboratorio 1 · Fundamentos del aislamiento acústico",
    2: "Laboratorio 2 · Modelos de predicción del aislamiento acústico",
}


@dataclass
class ContentBlock:
    kind: str
    text: str
    level: int = 0


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def _set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def _set_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr_text, fld_char2])


def _configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.0)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08

    for style_name, size, color in (
        ("Title", 24, "123B5D"),
        ("Heading 1", 17, "0B5FA5"),
        ("Heading 2", 14, "126782"),
        ("Heading 3", 12, "2B6777"),
    ):
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True

    if "Ecuación" not in styles:
        eq_style = styles.add_style("Ecuación", WD_STYLE_TYPE.PARAGRAPH)
        eq_style.font.name = "Cambria Math"
        eq_style.font.size = Pt(11.5)
        eq_style.paragraph_format.left_indent = Cm(0.7)
        eq_style.paragraph_format.space_before = Pt(5)
        eq_style.paragraph_format.space_after = Pt(7)

    header = section.header.paragraphs[0]
    header.text = "Diplomado en Acústica Aplicada a la Edificación"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(90, 103, 116)

    footer = section.footer.paragraphs[0]
    _set_page_number(footer)


def _clean_html(text: str) -> str:
    text = re.sub(r"<br\s*/?>", "\n", text, flags=re.I)
    text = re.sub(r"</?(?:b|strong|em|i|span|div|p|h\d|small|code)[^>]*>", "", text, flags=re.I)
    text = re.sub(r"<[^>]+>", "", text)
    replacements = {
        "&nbsp;": " ", "&amp;": "&", "&lt;": "<", "&gt;": ">",
        "&ge;": "≥", "&le;": "≤", "&rarr;": "→",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    return re.sub(r"\n{3,}", "\n\n", text).strip()


def _latex_to_linear(text: str) -> str:
    """Convierte LaTeX frecuente a una expresión lineal editable en Word."""
    text = text.strip().replace("\\,", " ").replace("\\;", " ")
    text = text.replace("\\qquad", "    ").replace("\\quad", "  ")
    text = re.sub(r"\\left|\\right", "", text)
    text = re.sub(r"\\mathrm\{([^{}]+)\}", r"\1", text)
    text = re.sub(r"\\text\{([^{}]+)\}", r"\1", text)
    text = re.sub(r"\\operatorname\{([^{}]+)\}", r"\1", text)
    symbols = {
        r"\\tau": "τ", r"\\alpha": "α", r"\\beta": "β", r"\\rho": "ρ",
        r"\\nu": "ν", r"\\theta": "θ", r"\\omega": "ω", r"\\pi": "π",
        r"\\Delta": "Δ", r"\\Sigma": "Σ", r"\\sum": "Σ", r"\\approx": "≈",
        r"\\geq": "≥", r"\\leq": "≤", r"\\cdot": "·", r"\\times": "×",
        r"\\infty": "∞", r"\\rightarrow": "→", r"\\pm": "±",
    }
    for old, new in symbols.items():
        text = text.replace(old, new)
    text = re.sub(r"\\log_\{10\}", "log₁₀", text)
    text = re.sub(r"\\log", "log", text)
    text = re.sub(r"\\sqrt\{([^{}]+)\}", r"√(\1)", text)
    # Fracciones simples; se repite para resolver fracciones anidadas sencillas.
    for _ in range(4):
        newer = re.sub(r"\\frac\{([^{}]+)\}\{([^{}]+)\}", r"(\1)/(\2)", text)
        if newer == text:
            break
        text = newer
    text = re.sub(r"_\{([^{}]+)\}", r"_\1", text)
    text = re.sub(r"\^\{([^{}]+)\}", r"^\1", text)
    text = text.replace("{", "").replace("}", "")
    text = re.sub(r"\\([A-Za-z]+)", r"\1", text)
    return re.sub(r"\s+", " ", text).strip()


def _literal(node: ast.AST) -> Any:
    try:
        return ast.literal_eval(node)
    except Exception:
        return None


def _call_name(call: ast.Call) -> str:
    parts: list[str] = []
    obj: ast.AST = call.func
    while isinstance(obj, ast.Attribute):
        parts.append(obj.attr)
        obj = obj.value
    if isinstance(obj, ast.Name):
        parts.append(obj.id)
    return ".".join(reversed(parts))


class _StageExtractor(ast.NodeVisitor):
    """Extrae contenido estático visible al alumno preservando el orden."""

    def __init__(self) -> None:
        self.blocks: list[ContentBlock] = []

    def _append(self, kind: str, value: Any, level: int = 0) -> None:
        if not isinstance(value, str):
            return
        value = _clean_html(value)
        if value:
            self.blocks.append(ContentBlock(kind, value, level))

    def visit_If(self, node: ast.If) -> Any:
        # No exportar ramas reservadas al docente. Cuando la condición pregunta
        # explícitamente por el rol Docente, solo se recorre la alternativa.
        try:
            condition = ast.unparse(node.test)
        except Exception:
            condition = ""
        if "Docente" in condition and ("role" in condition or "session_state" in condition):
            for statement in node.orelse:
                self.visit(statement)
            return
        for statement in node.body:
            self.visit(statement)
        for statement in node.orelse:
            self.visit(statement)

    def visit_Call(self, node: ast.Call) -> Any:
        name = _call_name(node)
        args = node.args

        if name == "header" and len(args) >= 3:
            for index, level in ((0, 1), (1, 2)):
                self._append("heading", _literal(args[index]), level)
            self._append("paragraph", _literal(args[2]))
        elif name in {"st.markdown", "st.write", "st.info", "st.warning", "st.caption", "st.error"} and args:
            value = _literal(args[0])
            if isinstance(value, str):
                cleaned = _clean_html(value)
                # Los títulos Markdown se transforman en encabezados Word.
                match = re.match(r"^(#{1,4})\s+(.+)$", cleaned)
                if match:
                    self._append("heading", match.group(2), min(3, len(match.group(1))))
                else:
                    self._append("paragraph", cleaned)
        elif name == "st.latex" and args:
            value = _literal(args[0])
            if isinstance(value, str):
                self.blocks.append(ContentBlock("equation", _latex_to_linear(value)))
        elif name in {"student_lesson", "lesson", "full_matter"}:
            # Estas funciones contienen normalmente título y cuerpo técnico.
            for arg in args:
                value = _literal(arg)
                if isinstance(value, str) and len(value.strip()) > 2:
                    self._append("paragraph", value)
        elif name in {"formative_development", "formative_numeric"}:
            # stage, key, pregunta, ...: solo se exporta el enunciado, no la pauta.
            if len(args) >= 3:
                self._append("question", _literal(args[2]))
        elif name == "check":
            # key, pregunta, opciones, respuesta correcta, retroalimentación.
            if len(args) >= 2:
                self._append("question", _literal(args[1]))
            if len(args) >= 3:
                options = _literal(args[2])
                if isinstance(options, (list, tuple)):
                    for option in options:
                        self._append("bullet", option)
        elif name in {"st.radio", "st.selectbox", "st.multiselect", "st.text_input", "st.text_area", "st.number_input"} and args:
            self._append("question", _literal(args[0]))

        self.generic_visit(node)


def _extract_stage_blocks(module_path: Path, stage: int) -> list[ContentBlock]:
    source = module_path.read_text(encoding="utf-8")
    tree = ast.parse(source)
    target = f"_stage{stage}_impl"
    for node in tree.body:
        if isinstance(node, (ast.FunctionDef, ast.AsyncFunctionDef)) and node.name == target:
            extractor = _StageExtractor()
            for statement in node.body:
                extractor.visit(statement)
            return extractor.blocks
    return []


def _add_cover(doc: Document, title: str, subtitle: str) -> None:
    doc.add_paragraph("DIPLOMADO EN ACÚSTICA APLICADA A LA EDIFICACIÓN", style="Subtitle")
    doc.add_paragraph(title, style="Title")
    p = doc.add_paragraph(subtitle)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(14)
    p.runs[0].font.color.rgb = RGBColor(60, 86, 110)
    doc.add_paragraph()
    note = doc.add_paragraph(
        "Apunte editable generado desde la vista del alumno. Los controles interactivos y las pautas reservadas al docente no se incluyen."
    )
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.runs[0].italic = True
    doc.add_page_break()


def _add_block(doc: Document, block: ContentBlock) -> None:
    if block.kind == "heading":
        doc.add_heading(block.text, level=max(1, min(3, block.level or 2)))
    elif block.kind == "equation":
        p = doc.add_paragraph(style="Ecuación")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(block.text)
    elif block.kind == "question":
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(7)
        run = p.add_run(block.text)
        run.bold = True
        run.font.color.rgb = RGBColor(18, 95, 130)
    elif block.kind == "bullet":
        doc.add_paragraph(block.text, style="List Bullet")
    else:
        for part in block.text.split("\n\n"):
            if part.strip():
                doc.add_paragraph(part.strip())


def build_laboratory_docx(project_root: str | Path, lab_number: int) -> bytes:
    """Genera un apunte Word editable con las etapas 0–10 del laboratorio."""
    root = Path(project_root)
    if lab_number not in (1, 2):
        raise ValueError("El laboratorio debe ser 1 o 2.")
    module_path = root / "labs" / f"laboratorio_{lab_number}.py"
    if not module_path.exists():
        raise FileNotFoundError(module_path)

    doc = Document()
    _configure_document(doc)
    _add_cover(doc, LAB_TITLES[lab_number], COURSE_TITLE)

    doc.add_heading("Contenido", level=1)
    for stage in range(11):
        doc.add_paragraph(f"Etapa {stage}", style="List Number")
    doc.add_page_break()

    for stage in range(11):
        doc.add_heading(f"Etapa {stage}", level=1)
        blocks = _extract_stage_blocks(module_path, stage)
        if not blocks:
            doc.add_paragraph("Contenido no disponible para exportación.")
        else:
            for block in blocks:
                _add_block(doc, block)
        if stage < 10:
            doc.add_page_break()

    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()


def _decode_answer(row: dict[str, Any] | None) -> dict[str, Any]:
    if not row:
        return {}
    payload = row.get("answer") or {}
    if isinstance(payload, dict) and "value" in payload and len(payload) == 1:
        try:
            payload = json.loads(payload["value"])
        except Exception:
            payload = {}
    return payload if isinstance(payload, dict) else {}


def _score(row: dict[str, Any] | None) -> float | None:
    if not row:
        return None
    value = row.get("teacher_score")
    if value is None:
        value = row.get("auto_score")
    return float(value or 0)


def _grade(percent: float) -> float:
    percent = max(0.0, min(100.0, float(percent)))
    return 1.0 + 3.0 * percent / 60.0 if percent < 60 else 4.0 + 3.0 * (percent - 60.0) / 40.0


def _add_summary_table(doc: Document, stage9: dict[str, Any] | None, stage10: dict[str, Any] | None) -> None:
    s9, s10 = _score(stage9), _score(stage10)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["Evaluación", "Puntaje", "Nota", "Estado"]
    for index, value in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = value
        _set_cell_shading(cell, "0B5FA5")
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.bold = True
    _set_repeat_table_header(table.rows[0])

    for label, row, maximum in (
        ("Laboratorio 2 · Etapa 9", stage9, 40),
        ("Laboratorio 2 · Etapa 10", stage10, 60),
    ):
        cells = table.add_row().cells
        value = _score(row)
        cells[0].text = label
        cells[1].text = "Pendiente" if value is None else f"{value:.1f}/{maximum}"
        cells[2].text = "—" if value is None else f"{_grade(value / maximum * 100):.1f}"
        cells[3].text = "Pendiente" if row is None else ("Revisada" if row.get("teacher_score") is not None else "Corrección automática")
        for cell in cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    total = None if s9 is None or s10 is None else s9 + s10
    p = doc.add_paragraph()
    run = p.add_run("Nota final del curso: ")
    run.bold = True
    p.add_run("Pendiente" if total is None else f"{_grade(total):.1f} ({total:.1f}/100 puntos)")


def _add_stage9_detail(doc: Document, row: dict[str, Any], questions: list[dict[str, Any]]) -> None:
    doc.add_heading("Laboratorio 2 · Etapa 9 · Evaluación de comprensión", level=1)
    payload = _decode_answer(row)
    answers = payload.get("answers", {}) if isinstance(payload, dict) else {}
    rubric = payload.get("rubric_scores", []) if isinstance(payload, dict) else []
    for index, item in enumerate(questions):
        doc.add_heading(f"Pregunta {index + 1} · {item.get('title', '')}", level=2)
        doc.add_paragraph(str(item.get("question", "")))
        chosen = answers.get(str(index)) if isinstance(answers, dict) else None
        correct_index = int(item.get("correct", 0))
        options = item.get("options", [])
        correct = options[correct_index] if 0 <= correct_index < len(options) else ""
        table = doc.add_table(rows=0, cols=2)
        table.style = "Table Grid"
        for label, value in (
            ("Respuesta del alumno", chosen or "Sin respuesta"),
            ("Respuesta correcta", correct),
            ("Corrección automática", item.get("explanation", "")),
            ("Puntaje", f"{float(rubric[index]) if index < len(rubric) else (4.0 if chosen == correct else 0.0):g}/4"),
        ):
            cells = table.add_row().cells
            cells[0].text, cells[1].text = label, str(value)
            _set_cell_shading(cells[0], "D9EAF7")
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Observación docente: ").bold = True
    p.add_run(str(row.get("teacher_note") or "Sin observación registrada."))


def _add_stage10_detail(doc: Document, row: dict[str, Any]) -> None:
    doc.add_heading("Laboratorio 2 · Etapa 10 · Diseño integrador", level=1)
    payload = _decode_answer(row)
    calculated = payload.get("calculated_result", {}) if isinstance(payload, dict) else {}
    student = payload.get("student_result", {}) if isinstance(payload, dict) else {}
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    data = [
        ("Resultado calculado", f"Rw(C; Ctr) = {calculated.get('rw', '—')} ({calculated.get('c', '—')}; {calculated.get('ctr', '—')}) dB"),
        ("Respuesta del alumno", f"Rw={student.get('rw', '—')} dB · C={student.get('c', '—')} dB · Ctr={student.get('ctr', '—')} dB"),
        ("Puntaje de diseño", f"{payload.get('design_score', 0):g}/40"),
        ("Puntaje de comprensión", f"{payload.get('comprehension_score', 0):g}/20"),
        ("Puntaje final", f"{_score(row) or 0:g}/60"),
    ]
    for label, value in data:
        cells = table.add_row().cells
        cells[0].text, cells[1].text = label, value
        _set_cell_shading(cells[0], "D9EAF7")
    for label, key in (("Muro/tabique", "wall"), ("Ventana", "window"), ("Puerta", "door")):
        data_item = payload.get(key, {}) if isinstance(payload, dict) else {}
        doc.add_heading(label, level=2)
        doc.add_paragraph(f"{data_item.get('description', 'Sin información')} · Rw {data_item.get('rw', '—')} dB")
    p = doc.add_paragraph()
    p.add_run("Observación docente: ").bold = True
    p.add_run(str(row.get("teacher_note") or "Sin observación registrada."))


def build_evaluation_docx(
    student_name: str,
    stage9_row: dict[str, Any] | None,
    stage10_row: dict[str, Any] | None,
    stage9_questions: list[dict[str, Any]] | None = None,
) -> bytes:
    """Genera un informe Word editable con notas, respuestas y correcciones."""
    doc = Document()
    _configure_document(doc)
    _add_cover(doc, "Informe de evaluaciones", COURSE_TITLE)
    doc.add_heading("Identificación", level=1)
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, value in (
        ("Alumno", student_name),
        ("Curso", COURSE_TITLE),
        ("Evaluaciones oficiales", "Laboratorio 2 · Etapas 9 y 10"),
    ):
        cells = table.add_row().cells
        cells[0].text, cells[1].text = label, value
        _set_cell_shading(cells[0], "D9EAF7")

    doc.add_heading("Resumen de notas", level=1)
    _add_summary_table(doc, stage9_row, stage10_row)
    if stage9_row:
        doc.add_page_break()
        _add_stage9_detail(doc, stage9_row, stage9_questions or [])
    if stage10_row:
        doc.add_page_break()
        _add_stage10_detail(doc, stage10_row)
    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()


def safe_filename(value: str) -> str:
    value = re.sub(r"[^A-Za-z0-9ÁÉÍÓÚáéíóúÑñ_-]+", "_", value.strip())
    return value.strip("_") or "alumno"


def build_evaluation_zip(
    grouped_rows: dict[str, dict[str, Any]],
    stage9_questions: list[dict[str, Any]] | None = None,
) -> bytes:
    output = io.BytesIO()
    with zipfile.ZipFile(output, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        for user_key, item in grouped_rows.items():
            name = str(item.get("name") or user_key)
            docx = build_evaluation_docx(
                name,
                item.get("stage9"),
                item.get("stage10"),
                stage9_questions,
            )
            archive.writestr(f"{safe_filename(name)}_evaluaciones.docx", docx)
    return output.getvalue()
